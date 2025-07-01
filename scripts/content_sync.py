#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import docx
import re
import os
import json
import hashlib
from datetime import datetime
from pathlib import Path

class ContentSync:
    def __init__(self, source_dir="public/documents", output_dir="src/content"):
        self.source_dir = source_dir
        self.output_dir = output_dir
        self.metadata_file = "content_metadata.json"
        
    def extract_word_content(self, file_path):
        """从Word文档提取内容"""
        try:
            doc = docx.Document(file_path)
            content = {
                'paragraphs': [],
                'tables': [],
                'sections': [],
                'metadata': {
                    'filename': os.path.basename(file_path),
                    'extracted_at': datetime.now().isoformat(),
                    'paragraph_count': len(doc.paragraphs),
                    'table_count': len(doc.tables)
                }
            }
            
            # 提取段落
            for i, para in enumerate(doc.paragraphs):
                if para.text.strip():
                    content['paragraphs'].append({
                        'index': i,
                        'text': para.text,
                        'style': para.style.name if para.style else 'Normal'
                    })
            
            # 提取表格
            for i, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                content['tables'].append({
                    'index': i,
                    'data': table_data
                })
            
            return content
            
        except Exception as e:
            print(f"读取文档时出错: {e}")
            return None
    
    def identify_sections(self, content):
        """识别文档章节结构"""
        sections = []
        current_section = None
        
        for para in content['paragraphs']:
            text = para['text'].strip()
            
            # 识别章节标题
            if re.match(r'^第[一二三四五六七八九十]+章', text) or \
               re.match(r'^\d+\.\d+', text) or \
               re.match(r'^[一二三四五六七八九十]+\.', text):
                
                if current_section:
                    sections.append(current_section)
                
                current_section = {
                    'title': text,
                    'level': self.get_section_level(text),
                    'content': [para],
                    'subsections': []
                }
            elif current_section:
                current_section['content'].append(para)
        
        if current_section:
            sections.append(current_section)
        
        return sections
    
    def get_section_level(self, title):
        """获取章节级别"""
        if re.match(r'^第[一二三四五六七八九十]+章', title):
            return 1
        elif re.match(r'^\d+\.\d+', title):
            return 2
        elif re.match(r'^[一二三四五六七八九十]+\.', title):
            return 2
        else:
            return 3
    
    def generate_markdown(self, sections, filename):
        """生成Markdown内容"""
        markdown_content = []
        
        # 添加文件头信息
        markdown_content.append("---")
        markdown_content.append(f"title: {filename}")
        markdown_content.append("lang: zh")
        markdown_content.append("layout: ../../layouts/Layout.astro")
        markdown_content.append("---")
        markdown_content.append("")
        
        # 添加导航
        markdown_content.append('<div class="top-nav">')
        markdown_content.append('  <a href="/">← 返回目录</a>')
        markdown_content.append('</div>')
        markdown_content.append("")
        
        # 生成章节内容
        for section in sections:
            level = section['level']
            title = section['title']
            
            # 添加章节标题
            if level == 1:
                markdown_content.append(f"# {title}")
            elif level == 2:
                markdown_content.append(f"## {title}")
            else:
                markdown_content.append(f"### {title}")
            
            markdown_content.append("")
            
            # 添加章节内容
            for para in section['content'][1:]:  # 跳过标题段落
                text = para['text'].strip()
                if text:
                    markdown_content.append(text)
                    markdown_content.append("")
        
        return "\n".join(markdown_content)
    
    def save_metadata(self, metadata):
        """保存元数据"""
        metadata_path = os.path.join(self.output_dir, self.metadata_file)
        os.makedirs(self.output_dir, exist_ok=True)
        
        with open(metadata_path, 'w', encoding='utf-8') as f:
            json.dump(metadata, f, ensure_ascii=False, indent=2)
    
    def get_file_hash(self, file_path):
        """获取文件哈希值用于检测变化"""
        with open(file_path, 'rb') as f:
            return hashlib.md5(f.read()).hexdigest()
    
    def check_for_updates(self):
        """检查源文件是否有更新"""
        metadata_path = os.path.join(self.output_dir, self.metadata_file)
        
        if not os.path.exists(metadata_path):
            return True  # 首次运行
        
        with open(metadata_path, 'r', encoding='utf-8') as f:
            old_metadata = json.load(f)
        
        updates_needed = []
        
        for filename in os.listdir(self.source_dir):
            if filename.endswith('.docx'):
                file_path = os.path.join(self.source_dir, filename)
                current_hash = self.get_file_hash(file_path)
                
                if filename not in old_metadata or \
                   old_metadata[filename]['hash'] != current_hash:
                    updates_needed.append(filename)
        
        return updates_needed
    
    def sync_content(self):
        """同步内容"""
        print("开始内容同步...")
        
        # 检查更新
        updates_needed = self.check_for_updates()
        
        if not updates_needed:
            print("所有文件都是最新的，无需更新")
            return
        
        print(f"需要更新的文件: {updates_needed}")
        
        metadata = {}
        
        for filename in os.listdir(self.source_dir):
            if filename.endswith('.docx'):
                file_path = os.path.join(self.source_dir, filename)
                
                print(f"处理文件: {filename}")
                
                # 提取内容
                content = self.extract_word_content(file_path)
                if not content:
                    continue
                
                # 识别章节
                sections = self.identify_sections(content)
                
                # 生成Markdown
                markdown_content = self.generate_markdown(sections, filename)
                
                # 保存Markdown文件
                output_filename = filename.replace('.docx', '.md')
                output_path = os.path.join(self.output_dir, output_filename)
                
                os.makedirs(self.output_dir, exist_ok=True)
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(markdown_content)
                
                # 记录元数据
                metadata[filename] = {
                    'hash': self.get_file_hash(file_path),
                    'last_updated': datetime.now().isoformat(),
                    'sections_count': len(sections),
                    'output_file': output_filename
                }
                
                print(f"已生成: {output_filename}")
        
        # 保存元数据
        self.save_metadata(metadata)
        print("内容同步完成！")

if __name__ == "__main__":
    sync = ContentSync()
    sync.sync_content() 
# -*- coding: utf-8 -*-

import docx
import re
import os
import json
import hashlib
from datetime import datetime
from pathlib import Path

class ContentSync:
    def __init__(self, source_dir="public/documents", output_dir="src/content"):
        self.source_dir = source_dir
        self.output_dir = output_dir
        self.metadata_file = "content_metadata.json"
        
    def extract_word_content(self, file_path):
        """从Word文档提取内容"""
        try:
            doc = docx.Document(file_path)
            content = {
                'paragraphs': [],
                'tables': [],
                'sections': [],
                'metadata': {
                    'filename': os.path.basename(file_path),
                    'extracted_at': datetime.now().isoformat(),
                    'paragraph_count': len(doc.paragraphs),
                    'table_count': len(doc.tables)
                }
            }
            
            # 提取段落
            for i, para in enumerate(doc.paragraphs):
                if para.text.strip():
                    content['paragraphs'].append({
                        'index': i,
                        'text': para.text,
                        'style': para.style.name if para.style else 'Normal'
                    })
            
            # 提取表格
            for i, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                content['tables'].append({
                    'index': i,
                    'data': table_data
                })
            
            return content
            
        except Exception as e:
            print(f"读取文档时出错: {e}")
            return None
    
    def identify_sections(self, content):
        """识别文档章节结构"""
        sections = []
        current_section = None
        
        for para in content['paragraphs']:
            text = para['text'].strip()
            
            # 识别章节标题
            if re.match(r'^第[一二三四五六七八九十]+章', text) or \
               re.match(r'^\d+\.\d+', text) or \
               re.match(r'^[一二三四五六七八九十]+\.', text):
                
                if current_section:
                    sections.append(current_section)
                
                current_section = {
                    'title': text,
                    'level': self.get_section_level(text),
                    'content': [para],
                    'subsections': []
                }
            elif current_section:
                current_section['content'].append(para)
        
        if current_section:
            sections.append(current_section)
        
        return sections
    
    def get_section_level(self, title):
        """获取章节级别"""
        if re.match(r'^第[一二三四五六七八九十]+章', title):
            return 1
        elif re.match(r'^\d+\.\d+', title):
            return 2
        elif re.match(r'^[一二三四五六七八九十]+\.', title):
            return 2
        else:
            return 3
    
    def generate_markdown(self, sections, filename):
        """生成Markdown内容"""
        markdown_content = []
        
        # 添加文件头信息
        markdown_content.append("---")
        markdown_content.append(f"title: {filename}")
        markdown_content.append("lang: zh")
        markdown_content.append("layout: ../../layouts/Layout.astro")
        markdown_content.append("---")
        markdown_content.append("")
        
        # 添加导航
        markdown_content.append('<div class="top-nav">')
        markdown_content.append('  <a href="/">← 返回目录</a>')
        markdown_content.append('</div>')
        markdown_content.append("")
        
        # 生成章节内容
        for section in sections:
            level = section['level']
            title = section['title']
            
            # 添加章节标题
            if level == 1:
                markdown_content.append(f"# {title}")
            elif level == 2:
                markdown_content.append(f"## {title}")
            else:
                markdown_content.append(f"### {title}")
            
            markdown_content.append("")
            
            # 添加章节内容
            for para in section['content'][1:]:  # 跳过标题段落
                text = para['text'].strip()
                if text:
                    markdown_content.append(text)
                    markdown_content.append("")
        
        return "\n".join(markdown_content)
    
    def save_metadata(self, metadata):
        """保存元数据"""
        metadata_path = os.path.join(self.output_dir, self.metadata_file)
        os.makedirs(self.output_dir, exist_ok=True)
        
        with open(metadata_path, 'w', encoding='utf-8') as f:
            json.dump(metadata, f, ensure_ascii=False, indent=2)
    
    def get_file_hash(self, file_path):
        """获取文件哈希值用于检测变化"""
        with open(file_path, 'rb') as f:
            return hashlib.md5(f.read()).hexdigest()
    
    def check_for_updates(self):
        """检查源文件是否有更新"""
        metadata_path = os.path.join(self.output_dir, self.metadata_file)
        
        if not os.path.exists(metadata_path):
            return True  # 首次运行
        
        with open(metadata_path, 'r', encoding='utf-8') as f:
            old_metadata = json.load(f)
        
        updates_needed = []
        
        for filename in os.listdir(self.source_dir):
            if filename.endswith('.docx'):
                file_path = os.path.join(self.source_dir, filename)
                current_hash = self.get_file_hash(file_path)
                
                if filename not in old_metadata or \
                   old_metadata[filename]['hash'] != current_hash:
                    updates_needed.append(filename)
        
        return updates_needed
    
    def sync_content(self):
        """同步内容"""
        print("开始内容同步...")
        
        # 检查更新
        updates_needed = self.check_for_updates()
        
        if not updates_needed:
            print("所有文件都是最新的，无需更新")
            return
        
        print(f"需要更新的文件: {updates_needed}")
        
        metadata = {}
        
        for filename in os.listdir(self.source_dir):
            if filename.endswith('.docx'):
                file_path = os.path.join(self.source_dir, filename)
                
                print(f"处理文件: {filename}")
                
                # 提取内容
                content = self.extract_word_content(file_path)
                if not content:
                    continue
                
                # 识别章节
                sections = self.identify_sections(content)
                
                # 生成Markdown
                markdown_content = self.generate_markdown(sections, filename)
                
                # 保存Markdown文件
                output_filename = filename.replace('.docx', '.md')
                output_path = os.path.join(self.output_dir, output_filename)
                
                os.makedirs(self.output_dir, exist_ok=True)
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(markdown_content)
                
                # 记录元数据
                metadata[filename] = {
                    'hash': self.get_file_hash(file_path),
                    'last_updated': datetime.now().isoformat(),
                    'sections_count': len(sections),
                    'output_file': output_filename
                }
                
                print(f"已生成: {output_filename}")
        
        # 保存元数据
        self.save_metadata(metadata)
        print("内容同步完成！")

if __name__ == "__main__":
    sync = ContentSync()
    sync.sync_content() 
 
 
 