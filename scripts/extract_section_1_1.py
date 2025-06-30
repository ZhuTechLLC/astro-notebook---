#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import docx
import re
import os

def extract_section_1_1(file_path):
    """专门提取1.1节的内容"""
    try:
        doc = docx.Document(file_path)
        print(f"文档段落数: {len(doc.paragraphs)}")
        
        # 查找1.1节
        section_1_1_content = []
        in_section_1_1 = False
        section_started = False
        
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            
            # 检测1.1节开始
            if "1.1" in text and ("自我觉知" in text or "投资智慧" in text):
                in_section_1_1 = True
                section_started = True
                section_1_1_content.append(f"# {text}")
                print(f"找到1.1节开始: {text}")
                continue
            
            # 检测1.1节结束（下一个主要章节）
            if in_section_1_1 and (text.startswith("1.2") or text.startswith("第二章")):
                print(f"1.1节结束，遇到: {text}")
                break
            
            # 收集1.1节内容
            if in_section_1_1 and text:
                section_1_1_content.append(text)
                print(f"添加内容: {text[:50]}...")
        
        # 如果没有找到内容，尝试更宽松的搜索
        if len(section_1_1_content) <= 1:
            print("尝试更宽松的搜索...")
            section_1_1_content = []
            in_section_1_1 = False
            
            for i, para in enumerate(doc.paragraphs):
                text = para.text.strip()
                
                # 检测1.1节开始（更宽松的条件）
                if "1.1" in text:
                    in_section_1_1 = True
                    section_1_1_content.append(f"# {text}")
                    print(f"找到1.1节开始: {text}")
                    continue
                
                # 检测1.1节结束
                if in_section_1_1 and (text.startswith("1.2") or text.startswith("第二章")):
                    print(f"1.1节结束，遇到: {text}")
                    break
                
                # 收集1.1节内容
                if in_section_1_1 and text:
                    section_1_1_content.append(text)
        
        # 生成Markdown文件
        markdown_content = """---
title: "1.1 自我觉知与投资智慧的起点"
lang: zh
layout: ../../layouts/Layout.astro
---

<div class="top-nav">
  <a href="/">← 返回目录</a>
  <a href="/000_Preface_CN">← 序言</a>
  <a href="/001_Chapter1_Know_Yourself_and_the_World_CN">← 返回第一章</a>
</div>

# 🧠 1.1 自我觉知与投资智慧的起点

"""
        
        # 添加内容
        for line in section_1_1_content[1:]:  # 跳过标题
            if line.strip():
                markdown_content += line + "\n\n"
        
        # 添加导航和样式
        markdown_content += """

---

<div class="bottom-nav">
  <a href="/000_Preface_CN">← 上一章：序言</a>
  <a href="/001_Chapter1/1.2_Understanding_the_World_CN">下一节：1.2 看懂世界 →</a>
</div>

<style>
  .example-box {
    background: linear-gradient(135deg, #e3f2fd 0%, #bbdefb 100%);
    border: 1px solid #2196f3;
    border-radius: 12px;
    padding: 1.5rem;
    margin: 1.5rem 0;
  }

  .tip-box {
    background: linear-gradient(135deg, #e8f5e8 0%, #c8e6c9 100%);
    border: 1px solid #4caf50;
    border-radius: 12px;
    padding: 1.5rem;
    margin: 1.5rem 0;
  }

  .warning-box {
    background: linear-gradient(135deg, #fff3e0 0%, #ffe0b2 100%);
    border: 1px solid #ff9800;
    border-radius: 12px;
    padding: 1.5rem;
    margin: 1.5rem 0;
  }
</style>
"""
        
        # 保存文件
        output_path = "../src/pages/001_Chapter1/1.1_Self_Awareness_and_Investment_Wisdom_CN.md"
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(markdown_content)
        
        print(f"✅ 1.1节内容已保存到: {output_path}")
        print(f"📝 提取了 {len(section_1_1_content)} 行内容")
        
        # 显示内容预览
        print("\n=== 内容预览 ===")
        for i, line in enumerate(section_1_1_content[:10]):
            print(f"{i+1}: {line}")
        
        return True
        
    except Exception as e:
        print(f"读取文档时出错: {e}")
        return False

if __name__ == "__main__":
    file_path = "../public/documents/美股投资实操手册I _06202025优化版本.docx"
    extract_section_1_1(file_path) 
# -*- coding: utf-8 -*-

import docx
import re
import os

def extract_section_1_1(file_path):
    """专门提取1.1节的内容"""
    try:
        doc = docx.Document(file_path)
        print(f"文档段落数: {len(doc.paragraphs)}")
        
        # 查找1.1节
        section_1_1_content = []
        in_section_1_1 = False
        section_started = False
        
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            
            # 检测1.1节开始
            if "1.1" in text and ("自我觉知" in text or "投资智慧" in text):
                in_section_1_1 = True
                section_started = True
                section_1_1_content.append(f"# {text}")
                print(f"找到1.1节开始: {text}")
                continue
            
            # 检测1.1节结束（下一个主要章节）
            if in_section_1_1 and (text.startswith("1.2") or text.startswith("第二章")):
                print(f"1.1节结束，遇到: {text}")
                break
            
            # 收集1.1节内容
            if in_section_1_1 and text:
                section_1_1_content.append(text)
                print(f"添加内容: {text[:50]}...")
        
        # 如果没有找到内容，尝试更宽松的搜索
        if len(section_1_1_content) <= 1:
            print("尝试更宽松的搜索...")
            section_1_1_content = []
            in_section_1_1 = False
            
            for i, para in enumerate(doc.paragraphs):
                text = para.text.strip()
                
                # 检测1.1节开始（更宽松的条件）
                if "1.1" in text:
                    in_section_1_1 = True
                    section_1_1_content.append(f"# {text}")
                    print(f"找到1.1节开始: {text}")
                    continue
                
                # 检测1.1节结束
                if in_section_1_1 and (text.startswith("1.2") or text.startswith("第二章")):
                    print(f"1.1节结束，遇到: {text}")
                    break
                
                # 收集1.1节内容
                if in_section_1_1 and text:
                    section_1_1_content.append(text)
        
        # 生成Markdown文件
        markdown_content = """---
title: "1.1 自我觉知与投资智慧的起点"
lang: zh
layout: ../../layouts/Layout.astro
---

<div class="top-nav">
  <a href="/">← 返回目录</a>
  <a href="/000_Preface_CN">← 序言</a>
  <a href="/001_Chapter1_Know_Yourself_and_the_World_CN">← 返回第一章</a>
</div>

# 🧠 1.1 自我觉知与投资智慧的起点

"""
        
        # 添加内容
        for line in section_1_1_content[1:]:  # 跳过标题
            if line.strip():
                markdown_content += line + "\n\n"
        
        # 添加导航和样式
        markdown_content += """

---

<div class="bottom-nav">
  <a href="/000_Preface_CN">← 上一章：序言</a>
  <a href="/001_Chapter1/1.2_Understanding_the_World_CN">下一节：1.2 看懂世界 →</a>
</div>

<style>
  .example-box {
    background: linear-gradient(135deg, #e3f2fd 0%, #bbdefb 100%);
    border: 1px solid #2196f3;
    border-radius: 12px;
    padding: 1.5rem;
    margin: 1.5rem 0;
  }

  .tip-box {
    background: linear-gradient(135deg, #e8f5e8 0%, #c8e6c9 100%);
    border: 1px solid #4caf50;
    border-radius: 12px;
    padding: 1.5rem;
    margin: 1.5rem 0;
  }

  .warning-box {
    background: linear-gradient(135deg, #fff3e0 0%, #ffe0b2 100%);
    border: 1px solid #ff9800;
    border-radius: 12px;
    padding: 1.5rem;
    margin: 1.5rem 0;
  }
</style>
"""
        
        # 保存文件
        output_path = "../src/pages/001_Chapter1/1.1_Self_Awareness_and_Investment_Wisdom_CN.md"
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(markdown_content)
        
        print(f"✅ 1.1节内容已保存到: {output_path}")
        print(f"📝 提取了 {len(section_1_1_content)} 行内容")
        
        # 显示内容预览
        print("\n=== 内容预览 ===")
        for i, line in enumerate(section_1_1_content[:10]):
            print(f"{i+1}: {line}")
        
        return True
        
    except Exception as e:
        print(f"读取文档时出错: {e}")
        return False

if __name__ == "__main__":
    file_path = "../public/documents/美股投资实操手册I _06202025优化版本.docx"
    extract_section_1_1(file_path) 