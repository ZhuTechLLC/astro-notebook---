#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import docx
import re

def read_word_document(file_path):
    """读取Word文档并提取内容"""
    try:
        doc = docx.Document(file_path)
        print(f"文档段落数: {len(doc.paragraphs)}")
        print(f"文档表格数: {len(doc.tables)}")
        
        # 提取所有文本
        full_text = []
        for i, para in enumerate(doc.paragraphs):
            if para.text.strip():
                full_text.append(f"段落{i+1}: {para.text}")
        
        # 查找附录部分
        appendix_sections = []
        in_appendix = False
        current_section = []
        
        for line in full_text:
            if "附录" in line or "Appendix" in line:
                if current_section:
                    appendix_sections.append("\n".join(current_section))
                current_section = [line]
                in_appendix = True
            elif in_appendix:
                current_section.append(line)
        
        if current_section:
            appendix_sections.append("\n".join(current_section))
        
        # 保存附录内容到文件
        with open("appendix_content.txt", "w", encoding="utf-8") as f:
            f.write("=== 附录部分内容 ===\n\n")
            for i, section in enumerate(appendix_sections):
                f.write(f"附录部分 {i+1}:\n")
                f.write(section)
                f.write("\n" + "="*50 + "\n\n")
        
        print("附录内容已保存到 appendix_content.txt")
        
        # 打印前几个段落作为预览
        print("\n=== 文档前10个段落预览 ===")
        for i in range(min(10, len(full_text))):
            print(full_text[i])
            
    except Exception as e:
        print(f"读取文档时出错: {e}")

if __name__ == "__main__":
    file_path = "public/documents/美股投资实操手册I _06202025优化版本.docx"
    read_word_document(file_path) 